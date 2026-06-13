"""Export session results to Word (.docx) format."""
from __future__ import annotations

import re
import tempfile
from pathlib import Path

from docx import Document

from models.schemas import IterationRecord, SessionResult


def _add_md(doc: Document, markdown: str) -> None:
    """Convert Markdown headings and paragraphs into a Word document."""
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            p.text = line[2:]
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            p.text = re.sub(r"^\d+\. ", "", line)
        elif line.strip() == "---":
            doc.add_paragraph("─" * 60)
        elif line.strip():
            doc.add_paragraph(line)


class DocxExporter:
    def export_final(self, result: SessionResult) -> Path:
        """Final document for a generator group as .docx."""
        doc = Document()
        doc.add_heading("AI Risk Council — Final Document", 0)

        meta = doc.add_paragraph()
        meta.add_run("Generator: ").bold = True
        meta.add_run(f"{result.config.generator_provider} / {result.config.generator_model}")

        score_p = doc.add_paragraph()
        score_p.add_run("Best Score: ").bold = True
        score_p.add_run(
            f"{result.best_score:.1f} / 100  |  "
            f"{len(result.iterations)} iteration(s)  |  "
            f"{result.termination_reason.replace('_', ' ').title()}"
        )
        doc.add_paragraph()

        _add_md(doc, result.final_policy.content)
        return _save(doc)

    def export_iteration_draft(self, result: SessionResult, iter_num: int) -> Path:
        """A single iteration's draft document as .docx."""
        rec = _get_iter(result, iter_num)
        doc = Document()
        doc.add_heading(f"AI Risk Council — Generator Draft (Iteration {iter_num})", 0)

        meta = doc.add_paragraph()
        meta.add_run("Generator: ").bold = True
        meta.add_run(f"{result.config.generator_provider} / {result.config.generator_model}")

        score_p = doc.add_paragraph()
        score_p.add_run("Score after review: ").bold = True
        score_p.add_run(f"{rec.aggregated_score:.1f} / 100")
        doc.add_paragraph()

        _add_md(doc, rec.policy_content)
        return _save(doc)

    def export_iteration_reviews(self, result: SessionResult, iter_num: int) -> Path:
        """All reviewer comments for one iteration as .docx."""
        rec = _get_iter(result, iter_num)
        doc = Document()
        doc.add_heading(f"AI Risk Council — Review Council Comments (Iteration {iter_num})", 0)

        agg = doc.add_paragraph()
        agg.add_run("Aggregated Score: ").bold = True
        agg.add_run(f"{rec.aggregated_score:.1f} / 100")
        doc.add_paragraph()

        for rv in rec.individual_reviews or []:
            doc.add_heading(f"Reviewer: {rv.reviewer_provider} / {rv.reviewer_model}", 2)

            if rv.score:
                sp = doc.add_paragraph()
                sp.add_run("Overall Score: ").bold = True
                sp.add_run(f"{rv.score.total_score:.1f} / 100")

                dims = doc.add_paragraph()
                for label, val in [
                    ("Fairness", rv.score.fairness),
                    ("Bias Mitigation", rv.score.bias_mitigation),
                    ("Ethical Soundness", rv.score.ethical_soundness),
                    ("Governance", rv.score.governance),
                    ("Controls", rv.score.controls),
                    ("Practicality", rv.score.practicality),
                ]:
                    dims.add_run(f"{label}: {val:.0f}   ")

            if rv.summary:
                doc.add_heading("Summary", 3)
                doc.add_paragraph(rv.summary)

            for title, items in [
                ("Biases Identified", rv.biases),
                ("Ethical Risks", rv.ethical_risks),
                ("Implementation Challenges", rv.implementation_challenges),
                ("Missing Controls", rv.missing_controls),
                ("Recommendations", rv.recommendations),
            ]:
                if items:
                    doc.add_heading(title, 3)
                    for item in items:
                        bp = doc.add_paragraph(style="List Bullet")
                        bp.text = item

            doc.add_paragraph()

        return _save(doc)


# ── Module-level helpers ──────────────────────────────────────────────────────

def _get_iter(result: SessionResult, iter_num: int) -> IterationRecord:
    rec = next((it for it in result.iterations if it.iteration_number == iter_num), None)
    if rec is None:
        raise ValueError(f"Iteration {iter_num} not found")
    return rec


def _save(doc: Document) -> Path:
    path = Path(tempfile.mktemp(suffix=".docx"))
    doc.save(str(path))
    return path
