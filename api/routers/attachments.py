"""File attachment upload and text extraction endpoint."""
from __future__ import annotations

import io
from pathlib import Path

from fastapi import APIRouter, HTTPException, UploadFile, File
from fastapi.responses import JSONResponse

router = APIRouter()

_MAX_BYTES = 10 * 1024 * 1024   # 10 MB

_ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".txt", ".md", ".csv"}


def _ext(filename: str) -> str:
    return Path(filename).suffix.lower()


def _parse_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n\n".join(parts).strip()


def _parse_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines).strip()


def _parse_xlsx(data: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _parse_txt(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return data.decode(enc).strip()
        except (UnicodeDecodeError, ValueError):
            continue
    return data.decode("latin-1", errors="replace").strip()


def _extract_text(filename: str, data: bytes) -> str:
    ext = _ext(filename)
    if ext == ".pdf":
        return _parse_pdf(data)
    if ext in (".docx", ".doc"):
        return _parse_docx(data)
    if ext in (".xlsx", ".xls"):
        return _parse_xlsx(data)
    if ext in (".txt", ".md", ".csv"):
        return _parse_txt(data)
    raise ValueError(f"Unsupported file type: {ext}")


@router.post("/parse")
async def parse_attachment(file: UploadFile = File(...)) -> JSONResponse:
    """Accept a document upload and return extracted plain text."""
    ext = _ext(file.filename or "")
    if ext not in _ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported file type '{ext}'. Allowed: {sorted(_ALLOWED_EXTENSIONS)}",
        )

    data = await file.read()
    if len(data) > _MAX_BYTES:
        raise HTTPException(status_code=413, detail="File exceeds 10 MB limit")

    try:
        text = _extract_text(file.filename or "", data)
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Could not parse file: {exc}") from exc

    if not text:
        raise HTTPException(status_code=422, detail="No readable text could be extracted from the file")

    return JSONResponse({
        "filename": file.filename,
        "content_type": file.content_type,
        "char_count": len(text),
        "word_count": len(text.split()),
        "text": text,
    })
